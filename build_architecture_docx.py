"""Generate Peekuu-Architecture.docx from the architecture overview."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PEACH = RGBColor(0xFF, 0x70, 0x43)
DARK = RGBColor(0x1F, 0x0B, 0x14)

doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PEACH
    return p

def h2(text):
    return doc.add_heading(text, level=2)

def para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    return t

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_heading('Peekuu — Tech Stack & Workflows', level=0)
for r in title.runs:
    r.font.color.rgb = DARK
para('Live-streaming + short-video commerce platform (Douyin-style feed x Bilibili-style '
     'danmaku) with two clients on one Firebase backend.', italic=True)

# ── 1. Tech stack ─────────────────────────────────────────────────────────────
h1('1. Complete Tech Stack')

h2('Mobile app (flutter_app/) — iOS + Android')
table(['Concern', 'Tech'], [
    ['Framework / language', 'Flutter, Dart'],
    ['State management', 'Riverpod (flutter_riverpod)'],
    ['Live RTC', 'agora_rtc_engine ^6.5.3 (pinned)'],
    ['Screen recording', 'flutter_screen_recording (Android MediaProjection)'],
    ['Video edit/encode', 'ffmpeg_kit_flutter_new (trim + thumbnail), video_player'],
    ['Captions (STT)', 'speech_to_text (device recognizer)'],
    ['Product scan', 'mobile_scanner (ML Kit barcode) + Claude Vision via backend'],
    ['Payments', 'in_app_purchase (App Store / Play), url_launcher (Stripe Connect)'],
    ['Firebase', 'core/auth/cloud_firestore/firebase_database/firebase_storage/cloud_functions'],
    ['UI / media', 'flutter_svg, google_fonts, cached_network_image, lottie, shimmer, flutter_animate'],
    ['Utilities', 'uuid, intl, path_provider, shared_preferences, permission_handler, image_picker, dio'],
    ['Icons / splash', 'flutter_launcher_icons, flutter_native_splash (P mark)'],
])

h2('Web app (web_app/) — Next.js creator studio + viewer')
table(['Concern', 'Tech'], [
    ['Framework', 'Next.js 16.2.4 (Turbopack), React 19, TypeScript'],
    ['Styling', 'Tailwind CSS'],
    ['State', 'Zustand (auth-store) + local hooks'],
    ['Live RTC', 'agora-rtc-sdk-ng ^4.24.3'],
    ['Virtual background (Room Mode)', 'agora-extension-virtual-background ^2.1.0 (WebGL/WASM segmentation)'],
    ['Clip recording', 'MediaRecorder + HTMLVideoElement.captureStream()'],
    ['Captions (STT)', 'Web Speech API'],
    ['Product scan', 'barcode scanner + frame-fingerprint + Claude Vision (/api/scan-product)'],
    ['Payments', 'Stripe Checkout (coin packs) via API routes + Stripe Connect (payouts)'],
    ['Firebase JS SDK', 'auth / firestore / database / storage / functions'],
])

h2('Backend (functions/) — Firebase Cloud Functions (Node 20, CommonJS, firebase-functions/v1)')
table(['Function', 'Purpose'], [
    ['getAgoraToken', 'Mint RTC tokens (dev mode = empty cert)'],
    ['validateApplePurchase / validateGooglePurchase', 'IAP receipt validation -> credit coins'],
    ['createConnectAccount / refreshConnectOnboardingLink', 'Stripe Connect onboarding'],
    ['requestDiamondPayout', 'Diamond->USD payout (tiered rate, debit-first, Stripe transfer)'],
    ['transcodeClipToMp4', 'Storage trigger: web .webm -> .mp4 + thumbnail (cross-platform playback)'],
    ['seedGiftCatalog', 'Seed giftCatalog/*'],
    ['cleanupOldStreams / trackPeakViewers', 'Housekeeping (pubsub schedule / Firestore trigger)'],
    ['Libraries', 'stripe, agora-token, googleapis, ffmpeg-static'],
])

h2('Data & infra (Firebase project qaramia-4c405)')
bullet('Firestore — users (+ wallet, creatorBalance, iapTransactions, payouts, following/followers, '
       'cohost_invites), videos (+ likes, comments), streams (+ gifts, cohosts, gifters), '
       'sponsorships (+ sends), giftCatalog, ads')
bullet('Realtime Database — danmaku/{streamId} (high-freq chat), captions/{streamId}/current (live transcript)')
bullet('Cloud Storage — videos/{id}.mp4|.webm + _thumb.jpg')
bullet('Auth — Email/Password + Google')
bullet('Security — firestore.rules, database.rules.json, storage.rules')

h2('Third-party services')
bullet('Agora.io (RTC) - Firebase/GCP - Stripe (Connect + Checkout) - Apple/Google IAP - '
       'Anthropic Claude (Vision product detection) - Google ML Kit (mobile barcode) - '
       'device/Web Speech (captions)')

h2('Bundle identity')
bullet('com.qaramia.app (unchanged) - domain qaramia.com (Stripe URLs) - display name "peekuu".')

# ── 2. Workflows ──────────────────────────────────────────────────────────────
h1('2. Workflows — input / output / task per app')
para('Both clients are dual-role (broadcast AND watch) and share the backend. The backend is the '
     'source of truth; clients are I/O + media.', italic=True)

h2('A. Auth')
bullet('Input: email/password or Google (both apps).')
bullet('Task: Firebase Auth; create/read users/{uid}.')
bullet('Output: session; profile.')

h2('B. Go Live (broadcast)')
table(['', 'Mobile', 'Web (Studio)'], [
    ['Input', 'camera + mic, title/category, Room Mode + bg, gift goal, filter', 'same, via Studio form'],
    ['Task', 'getAgoraToken -> join Agora as broadcaster; create streams/{id}; apply beauty/VB; screen-record (MediaProjection)',
     'join Agora as host; create streams/{id}; VB via extension; MediaRecorder clip capture'],
    ['Output', 'live A/V to Agora; streams/{id} metadata; recorded mp4', 'live A/V; streams/{id}; recorded .webm'],
])

h2('C. Watch a stream')
bullet('Input: tap a live stream; chat text; gift taps.')
bullet('Task: join Agora as audience; subscribe to host (+ co-host); onChildAdded on danmaku/{streamId}; '
       'render captions from captions/.../current.')
bullet('Output: rendered video, bullet-comment overlay, gift animations, product drawer, leaderboard, gift-goal bar.')

h2('D. Danmaku (chat)')
bullet('Input: message text. Task: push to RTDB danmaku/{streamId} (validated authorUid==auth.uid). '
       'Output: real-time scroll for all viewers.')

h2('E. Gifting (monetization core)')
bullet('Input: select gift (tiered picker; sponsored row when applicable).')
bullet('Task: atomic Firestore txn — debit wallet.coins, credit creatorBalance.diamonds, write '
       'streams/{id}/gifts/{id}, upsert streams/{id}/gifters/{uid} (leaderboard), bump streams.totalGifts; '
       'sponsored sends also record sponsorships/{id}/sends.')
bullet('Output: gift animation, updated leaderboard + gift-goal progress, danmaku "sent gift".')

h2('F. Product detection + ads')
bullet('Input: stream frames + host speech.')
bullet('Task: barcode (ML Kit/web) -> lookup; else Claude Vision (/api/scan-product); spoken-product '
       'detection from caption transcript; matchAd.')
bullet('Output: streams/{id}.featuredProducts + featuredAd -> floating bag + slide-up ProductDrawer; '
       'session earnings estimate.')

h2('G. Captions')
bullet('Input: host mic (Web Speech API / device STT). Task: transcribe, throttle-write '
       'captions/{streamId}/current. Output: viewer caption overlay. NOTE: web works live; mobile CC is '
       'paused while broadcasting (Agora holds the mic).')

h2('H. Co-host')
bullet('Input: host invites a viewer -> viewer accepts.')
bullet('Task: invitee role-switches to broadcaster, publishes; both publish to one Agora channel.')
bullet('Output: 50/50 side-by-side split on host + all viewers (both platforms).')

h2('I. Room Mode (virtual background)')
bullet('Input: toggle + room preset image. Task: Agora VB extension/native segmentation composites the '
       'room photo behind the host. Output: composited broadcast. NOTE: web VB only runs in a production '
       'build, not the Turbopack dev server.')

h2('J. Post-stream clip -> feed')
table(['', 'Mobile', 'Web'], [
    ['Input', 'recorded mp4', 'recorded webm'],
    ['Task', 'Post-stream editor: trim (FFmpeg cut), zoom/blur/vignette/text/stickers/filter (metadata), '
     'caption -> VideoUploadService upload + videos/{id}',
     'PostStreamEditor: same effects as metadata, trim = non-destructive window, caption -> publishClip '
     'upload + videos/{id}; transcodeClipToMp4 then converts webm->mp4 + thumbnail'],
    ['Output', 'feed video (mp4) with effects', 'feed video (mp4 after transcode)'],
])

h2('K. Short-video feed')
bullet('Input: scroll. Task: read videos (desc), play with video_player/<video>, apply effects at playback '
       '(composeVideo / ComposedVideo: vignette->blur->filter->zoom->text+stickers, honoring trim window). '
       'Output: vertical feed; like/comment/share.')

h2('L. Money flow')
bullet('Coins in: Mobile = IAP -> validate*Purchase credits wallet; Web = Stripe Checkout -> webhook credits wallet.')
bullet('Diamonds: earned by creators from gifts (creatorBalance).')
bullet('Payout out: Stripe Connect onboarding -> requestDiamondPayout (tier rate: Rising $0.010 / '
       'Partner $0.012 / Elite $0.014; min 5,000 diamonds) -> Stripe transfer -> payouts/{id} ledger.')

h2('Designed, not built')
bullet('Watch-to-Earn (viewer pay-for-attention: heartbeats, daily caps, anti-fraud, earn-coin cash-out) — '
       'spec only in docs/Watch-to-Earn-Design.md.')

doc.save('E:/claude/qaramia-v2/Peekuu-Architecture.docx')
print('Wrote E:/claude/qaramia-v2/Peekuu-Architecture.docx')
