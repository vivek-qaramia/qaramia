"""Generate Peekuu-Competitive-Analysis.docx — verified competitive analysis
of the 9 key US + Asian live-commerce players, the opportunity gap, and
Peekuu's differentiation. Figures are fact-checked (deep-research pass);
confidence + sources are footnoted at the end."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PEACH = RGBColor(0xFF, 0x70, 0x43)
GOLD = RGBColor(0xC9, 0x7A, 0x00)
WINE = RGBColor(0x1F, 0x0B, 0x14)
GREY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x1B, 0x7F, 0x4B)
RED = RGBColor(0xB0, 0x2A, 0x37)

doc = Document()
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = PEACH
    return p


def h2(text):
    return doc.add_heading(text, level=2)


def para(text, italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    return p


def conf(label):
    """Inline confidence tag."""
    return f'  [{label}]'


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()
    return t


# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading('peekuu — Competitive Analysis', level=0)
for r in title.runs:
    r.font.color.rgb = WINE
para('US & Asian live-streaming + short-video commerce — the 9 key players, the opportunity '
     'gap, and how peekuu differentiates.', italic=True)
para('Figures fact-checked via a multi-source research pass (24 sources, 25 claims adversarially '
     'verified). Confidence is tagged inline: [High] = primary/3-0 verified, [Med] = 2-1 vote, '
     '[Low] = blog-only / unverified. Numbered superscripts map to Sources at the end. Claims that '
     'FAILED verification are listed under "Do not cite" so they never reach a deck.',
     italic=True, color=GREY, size=9)

# ── 0. Headline ────────────────────────────────────────────────────────────────
h1('The headline: a ~20x scale gap and an AI gap nobody has closed')
para('Asia runs a mature, closed-loop live-commerce model (feed -> live -> shop -> pay in one '
     'app); the US is early and fragmented. The gap is structural, not a demand problem.')
table(['Market', 'Live-commerce size', 'Confidence'], [
    ['China', '~$695B (2023) -> ~$703B (2024, eMarketer); >$1T by 2026', 'High [1]'],
    ['Asia overall', '~$370B (2024, segment-scoped)', 'High [1]'],
    ['United States', '~$17B', 'High [1]'],
])
para('China figures are segment-scoped (beauty/fashion/food/electronics); broader GMV estimates '
     'run higher (~$807B).', italic=True, color=GREY, size=9)

# ── 1. Features ─────────────────────────────────────────────────────────────────
h1('1. Features')
table(['Platform', 'Feed', 'Live', 'Danmaku', 'Co-host', 'In-stream shop', 'Post-edit'], [
    ['peekuu', 'Yes', 'Yes', 'Yes (native)', 'Yes 50/50', 'Yes (auto-detect)', 'Yes (full)'],
    ['Douyin', 'Yes', 'Yes', 'Yes', 'Yes (PK)', 'Yes (closed-loop)', 'Yes'],
    ['Taobao Live', 'Weak', 'Yes', 'Yes', 'Yes', 'Yes (marketplace)', 'Limited'],
    ['Kuaishou', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
    ['Bilibili', 'Long-form', 'Yes', 'Yes (origin)', 'Limited', 'Limited', 'Yes'],
    ['TikTok / Shop', 'Best-in-class', 'Yes', 'Side chat', 'Yes', 'Yes (manual tag)', 'Yes'],
    ['Whatnot', 'No', 'Yes', 'Side chat', 'Limited', 'Yes (auctions)', 'No'],
    ['Twitch', 'No', 'Yes', 'Side chat', 'Yes (Guest Star)', 'No', 'Clips'],
    ['YouTube', 'Shorts', 'Yes', 'No', 'Limited', 'Yes (manual)', 'Yes'],
])
bullet('the feed-first model is winning: by 2022 GMV, Douyin 47% and Kuaishou 27% vs. '
       'marketplace-native Taobao trailing both — despite Taobao keeping ~74% user adoption. '
       'Content out-converts the store, which is exactly peekuu’s architecture. [High] [2]',
       bold_lead='Verified strategic signal: ')
bullet('the only Western Asian-style fusion (organic feed + live + shop); SE Asia GMV '
       '~$16.3B (2023) -> ~$25-30B (2024), #2 ahead of Lazada. [Med, conflates total vs. '
       'live-only GMV] [3]', bold_lead='TikTok Shop: ')

# ── 2. AI ────────────────────────────────────────────────────────────────────────
h1('2. AI usage — peekuu’s genuine white space')
para('The most important verified finding: automatic real-time multimodal vision+speech product '
     'detection is technically proven but shipped by NO US incumbent.', bold=True)
bullet('"Livestreaming Product Retrieval" (LPR), arXiv 2407.16248, ACM MM ’24 — the SGMN '
       'model uses the salesperson’s spoken words (via ASR) plus video to auto-identify the '
       'on-screen product (R@1 38.7% -> 43.4% with the speech module). This is peekuu’s exact '
       'premise, peer-reviewed. [High] [4]', bold_lead='Academic proof of the premise: ')
bullet('its deployed "Identify Similar Objects" visual search (2024 US/UK pilot) is vision-only, '
       'post-hoc (mostly recorded video), and viewer-initiated/opt-in — NOT real-time, NOT '
       'speech-fused, NOT barcode. [High] [5]', bold_lead='TikTok’s shipped AI is weaker: ')
para('Where incumbents lead: recommendation (TikTok best-in-class; Taobao personalization). '
     'peekuu is behind there — its AI leads on attribution/shoppability, not feed-ranking. Be '
     'honest about this split in any pitch.', italic=True)

# ── 3. Monetization ──────────────────────────────────────────────────────────────
h1('3. Monetization / commission take rates')
table(['Platform', 'Shop commission', 'Confidence'], [
    ['TikTok Shop (US)', '5-6% referral, no separate txn fee (+20%-of-referral refund admin fee)', 'High [6]'],
    ['Whatnot', '8% standard (4% Coins/Money, 5% Electronics) + 2.9% + $0.30 processing', 'High [7]'],
    ['Douyin / Kuaishou / Taobao', 'Closed-loop ad + commission (not claim-verified)', '-'],
    ['Shopee Live', 'Per-country; precise ranges FAILED verification — do not cite', 'Low'],
])
bullet('Whatnot is the leading standalone US live-shopping marketplace: >$2B livestream sales '
       'Jan-Sep 2024 (a Jan 2026 promo adds 0% commission above $1,500 in select categories). '
       '[High] [7]')

# ── 4. Gift giving ───────────────────────────────────────────────────────────────
h1('4. Gift giving — the fairness benchmark is the story')
table(['Platform', 'Creator economics', 'Confidence'], [
    ['TikTok gifting', 'coin ~$0.01 in, diamond ~$0.005 out (2:1); nominal ~50% but EFFECTIVE '
     'platform cut ~50-77% (ABC: only ~40% reached creator; FXC: 77%)', 'High [8]'],
    ['Twitch', '50/50 default subs; 70/30 uncapped since Jan 2024 (Tier 1 $4.99 -> ~$2.50)', 'High [9]'],
    ['YouTube Super Chat', '70% to creator ($10 browser -> ~$7; iOS -> ~$4.90 after Apple)', 'High [10]'],
    ['Douyin gifting', '"Billion-dollar" economy but split not claim-verified', 'Low'],
])
bullet('TikTok — the platform with the most live-commerce traction — has the WORST creator '
       'split (creators net as little as 40%). Twitch (70/30) and YouTube (70%) set the fairness '
       'benchmark. peekuu’s tiered cash-out (Rising $0.010 / Partner $0.012 / Elite $0.014 '
       'per diamond) is positioned to beat TikTok and approach Twitch — a concrete '
       'creator-acquisition wedge.', bold_lead='Takeaway: ')
bullet('all splits worsen ~30% on mobile due to Apple/Google IAP fees — applies to peekuu too. '
       'Twitch/YouTube subscriptions + gifted subs (recurring revenue) are a gap peekuu has not '
       'yet filled.', bold_lead='Honesty caveat: ')

# ── 5. Gamification ──────────────────────────────────────────────────────────────
h1('5. Gamification')
para('Thinnest-verified dimension (mostly blog sources) — treat as directional.', italic=True,
     color=GREY)
bullet('Asian PK battles (Douyin/Kuaishou creator-vs-creator gift races) are widely reported as '
       'the highest-grossing live mechanic — figures unverified. [Low]')
bullet('Twitch is the Western gamification benchmark: channel points, Hype Train (escalation), '
       'loyalty currency. [Low, well-established]')
bullet('peekuu today has leaderboards + gift goals; it LACKS PK battles, channel points/loyalty '
       'currency, streaks, Hype Train-style escalation, and watch-to-earn.')
bullet('peekuu already has co-host + gift goals + leaderboards, so a creator-vs-creator PK '
       'battle is mostly assembling existing plumbing — and no US incumbent has copied it.',
       bold_lead='Highest-leverage build: ')

# ── 6. Opportunity gap ───────────────────────────────────────────────────────────
h1('6. The opportunity gap peekuu is addressing')
para('The West has a ~$17B live-commerce market against Asia’s ~$370B (~20x). The reason '
     'is not demand — it is that no Western platform has fused the content feed, the live room, '
     'and automatic shoppability into one loop the way Douyin did.', bold=True)

h2('Gap 1 — The fused-product gap (structural)')
para('In the US the four ingredients of Asian live commerce are split across four apps:')
table(['What’s needed', 'Who has it', 'Who’s missing it'], [
    ['Short-video discovery feed', 'TikTok, YouTube', 'Whatnot, Twitch'],
    ['Live + community (danmaku)', 'Twitch, Bilibili', 'TikTok, Whatnot'],
    ['Native in-stream shopping', 'Whatnot, TikTok Shop', 'Twitch, YouTube'],
    ['Creator gifting + payouts', 'Twitch, YouTube', 'Whatnot'],
])
bullet('peekuu is the only Western product carrying all four columns at once. The China data '
       '(Douyin > Taobao) proves the fused, feed-first architecture is what wins. [High] [2]')

h2('Gap 2 — The AI attribution gap (the real moat)')
bullet('Every Western live-shopping experience requires manual product pre-tagging — friction, '
       'and TikTok Shop’s known onboarding bottleneck.')
bullet('peekuu removes it: vision + speech + barcode, ~2s, zero creator effort — proven in '
       'research, shipped by no US incumbent. [High] [4][5]')
bullet('Doing so generates a proprietary multimodal corpus (catalog + transcripts + advertiser '
       'performance) that compounds with usage — a data network effect a UI clone cannot copy.')

h2('Gap 3 — The creator-economics gap (the recruiting wedge)')
bullet('The platform with the most traction (TikTok) pays the least (~40% to creators); '
       'Twitch/YouTube pay ~70%. peekuu’s tiered split is a direct reason for a mid-tier '
       'creator to switch. [High] [8][9][10]')

h2('Gap 4 — The engagement-mechanics gap (the fast follow)')
bullet('Asia’s highest-grossing live mechanic (PK battles) and watch-to-earn have no US '
       'equivalent; peekuu already has the rails (co-host, gift goals, leaderboards) to ship '
       'them. [Low/directional]')

# ── 7. Differentiation ───────────────────────────────────────────────────────────
h1('7. How peekuu differentiates — one table')
table(['Dimension', 'The market’s gap', 'peekuu’s differentiation'], [
    ['Structure', 'Feed, live, shop, gifting split across 4 apps', 'One fused loop (the Douyin model, localized)'],
    ['AI', 'Manual tagging everywhere; TikTok pilot is vision-only/opt-in', 'Automatic vision+speech+barcode detection, ~2s — shipped by no one'],
    ['Data', 'UI clones are easy', 'Compounding multimodal corpus (catalog + transcripts + ad performance)'],
    ['Creator pay', 'TikTok pays as little as 40%', 'Tiered cash-out beating TikTok, approaching Twitch’s 70%'],
    ['Engagement', 'US lacks Asia’s PK / escalation loops', 'PK battles + watch-to-earn on existing co-host/gifting rails'],
    ['Compliance', 'Captions a cost', 'Accessibility (ADA / EU Act) shipped as a feature'],
])
para('The defensible core: peekuu is the Western fusion of Asia’s proven feed-first '
     'live-commerce loop, with the one piece even Asia has not productized — automatic multimodal '
     'shoppability — which compounds into a data moat while paying creators more than the '
     'incumbent that currently owns the category.', bold=True, color=WINE)
para('The honest counter (be ready for it): the gap is real and verified, but liquidity / '
     'cold-start (a two-sided network is worthless empty) and TikTok bolting on auto-detection '
     'are the two risks that decide whether peekuu fills the gap first. The defense to both is '
     'the same — creator-earnings lock-in + the proprietary detection corpus, neither clonable '
     'by copying features.', italic=True)

# ── Do not cite ──────────────────────────────────────────────────────────────────
h1('Do not cite — claims that FAILED verification')
for c in [
    'US live commerce reaching ~$680B by 2030 (killed 0-3).',
    'Shopee 2024 SE Asia GMV ~$83.4B (killed 0-3).',
    'TikTok Shop US GMV $17.5B in 2024 (killed 1-2).',
    'TikTok crossing $6B in-app-purchase revenue in 2024 (killed 1-2).',
    'Precise Shopee / TikTok SE Asia per-country commission ranges (killed 1-2).',
]:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(c)
    r.font.color.rgb = RED

# ── Coverage gaps ────────────────────────────────────────────────────────────────
h1('Coverage gaps (verify before a high-stakes slide)')
for c in [
    'Douyin / Kuaishou / Bilibili native gift splits and coin economics.',
    'Shopee Live gamification mechanics and Taobao’s AI stack — not claim-verified.',
    'Live-commerce-specific GMV (vs. total platform GMV) for Shopee Live, TikTok Shop, Whatnot.',
]:
    bullet(c)

# ── Sources ──────────────────────────────────────────────────────────────────────
h1('Sources')
sources = [
    'Live-commerce market size (China / Asia / US): ecdb.com; modernretail.co; electroiq.com.',
    'China GMV share — Douyin 47% / Kuaishou 27% / Taobao trailing (2022): ecdb.com (Statista 1339406).',
    'TikTok Shop SE Asia GMV 2023-2024 (#2 ahead of Lazada): sellercraft.co (Momentum Works "Ecommerce in SE Asia 2025").',
    'Livestreaming Product Retrieval / SGMN vision+speech model: arXiv 2407.16248 (ACM MM ’24).',
    'TikTok "Identify Similar Objects" visual-search pilot (vision-only, opt-in): business-humanrights.org.',
    'TikTok Shop US take rate 5-6%: cube.asia TikTok Shop take-rate tracker (TikTok US Seller Academy, May 2025).',
    'Whatnot fees (8% + 2.9%+$0.30) and >$2B Jan-Sep 2024: help.whatnot.com; modernretail.co.',
    'TikTok gifting economics (coin/diamond; ~50-77% effective cut): naavik.co; influencermarketinghub.com; FXC Intelligence; ABC investigation.',
    'Twitch splits (50/50; 70/30 uncapped Jan 2024): blog.twitch.tv (Jan 24 2024); influencermarketinghub.com.',
    'YouTube Super Chat 70% creator share: support.google.com/youtube; influencermarketinghub.com.',
]
for i, s in enumerate(sources, start=1):
    p = doc.add_paragraph()
    r = p.add_run(f'[{i}] ')
    r.bold = True
    r.font.size = Pt(9)
    r2 = p.add_run(s)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY

doc.save('E:/claude/qaramia-v2/Peekuu-Competitive-Analysis.docx')
print('Wrote E:/claude/qaramia-v2/Peekuu-Competitive-Analysis.docx')
