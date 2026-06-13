"""Generate two tailored Peekuu pitch decks (Founders Fund + Benchmark cuts)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK

PEACH = RGBColor(0xFF, 0x70, 0x43)
GOLD = RGBColor(0xC9, 0x7A, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)


def new_doc():
    doc = Document()
    n = doc.styles['Normal']
    n.font.name = 'Calibri'
    n.font.size = Pt(11)
    return doc


_counter = {'n': 0}


def slide(doc, num, total, title, goal, bullets, first=False):
    # Auto-number so inserting a slide doesn't require renumbering every call.
    # The positional `num` is ignored; the counter resets on the first slide.
    if first:
        _counter['n'] = 0
    _counter['n'] += 1
    num = _counter['n']
    if not first:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # kicker
    k = doc.add_paragraph()
    kr = k.add_run(f'SLIDE {num} / {total}')
    kr.font.size = Pt(9)
    kr.font.bold = True
    kr.font.color.rgb = GREY
    # title
    h = doc.add_heading(title, level=1)
    for r in h.runs:
        r.font.color.rgb = PEACH
    # goal (what this slide must prove to THIS firm)
    g = doc.add_paragraph()
    gr = g.add_run('What this slide must land: ')
    gr.bold = True
    gr.font.color.rgb = GOLD
    gr.font.size = Pt(10)
    gr2 = g.add_run(goal)
    gr2.italic = True
    gr2.font.size = Pt(10)
    gr2.font.color.rgb = GOLD
    # bullets
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')


def cover(doc, fund, lens):
    t = doc.add_heading('peekuu', level=0)
    for r in t.runs:
        r.font.color.rgb = PEACH
    sub = doc.add_paragraph()
    sr = sub.add_run('Live Streaming & Commerce — the camera connected to the checkout')
    sr.italic = True
    sr.font.size = Pt(13)
    doc.add_paragraph()
    f = doc.add_paragraph()
    fr = f.add_run(f'Pitch cut for: {fund}')
    fr.bold = True
    fr.font.size = Pt(12)
    n = doc.add_paragraph()
    nr = n.add_run(f'Framing lens: {lens}')
    nr.font.size = Pt(10)
    nr.font.color.rgb = GREY
    note = doc.add_paragraph()
    notr = note.add_run('Numbers in [brackets] are placeholders — insert your real figures before sending. '
                        'The gold "What this slide must land" line is a coaching note; delete it in the final deck.')
    notr.italic = True
    notr.font.size = Pt(9)
    notr.font.color.rgb = GREY


# ── FOUNDERS FUND CUT — vision / secret / monopoly first ──────────────────────
ff = new_doc()
cover(ff, 'Founders Fund', 'Contrarian secret -> proprietary tech -> monopoly (Zero to One). Vision-led; metrics secondary.')
T = 12
slide(ff, 1, T, 'The one-liner',
      'A contrarian, definite belief stated as fact. Make the partner lean in.',
      ['Every product shown OR spoken in a live stream becomes a shoppable ad in real time.',
       'Western platforms have live video + short video — nobody has connected the camera to the checkout.',
       'peekuu turns the content itself into ad inventory.'], first=True)
slide(ff, 2, T, 'The secret',
      'The non-obvious truth others miss + why it is hard (so it is defensible).',
      ['Live commerce in the West fails on attribution: matching the right product/ad to the moment is a hard multimodal problem.',
       'We solved it: barcode + Claude Vision (what is seen) + speech/caption detection (what is said) -> product + targeted ad in ~2 seconds.',
       'This is an AI problem that only became cheap to solve now — that is the unfair timing.'])
slide(ff, 3, T, 'Why now',
      'Inevitability — the wave is forming and we are early on it.',
      ['AI made real-time multimodal detection cheap and accurate.',
       'TikTok Shop proved Western appetite for live/social commerce (and left the attribution layer open).',
       'Regulatory tailwind: EU Accessibility Act + ADA make captions mandatory — we ship them as a feature, not a cost.'])
slide(ff, 4, T, 'The opportunity gap — 20x and wide open',
      'Why the white space exists and why it is defensible. (Figures fact-checked.)',
      ['Asia live commerce is ~$370B; the US is ~$17B — a ~20x gap that is structural, not a demand problem.',
       'In China the feed-first model won: Douyin took 47% of live-commerce GMV while marketplace-native Taobao trailed — content out-converts the store.',
       'In the West the four pieces — feed, live, danmaku, shoppability — sit in four separate apps (TikTok, Twitch, Whatnot, YouTube each own one). peekuu fuses all four.',
       'The piece even Asia has not productized: automatic vision+speech product detection. TikTok\'s pilot is vision-only and opt-in — peekuu ships the full auto-detect.'])
slide(ff, 5, T, 'The magic moment (live demo)',
      'One demo that makes the secret undeniable. Do this live, not on a slide.',
      ['A streamer says "I love this Stanley cup" — 2 seconds later a shoppable product card + matched ad appears for every viewer.',
       'No tagging, no manual setup. Vision + voice did it automatically.',
       'This single demo is the pitch.'])
slide(ff, 5, T, 'Monopoly thesis — the compounding moat',
      'Why this becomes a monopoly, not a feature a giant clones.',
      ['Flywheel: more streams -> larger product catalog + transcript corpus + advertiser network -> detection accuracy improves on BOTH signals -> better targeting -> more advertisers + creators.',
       'The moat is the proprietary multimodal dataset, not the UI. Cloning the interface does not clone the data advantage.',
       'Switching costs accrue: creator earnings + audience, advertiser performance history.'])
slide(ff, 6, T, 'The second act is the real prize',
      'The 10x-bigger business hiding behind the consumer app.',
      ['Brand-intelligence SaaS: "show me every mention of my brand or competitor across peekuu last week."',
       'Comparables (Sprinklr, Talkwalker) charge $20K-$200K/year — pure-margin SaaS on top of the ad model.',
       'Live commerce is the consumer wedge; the multimodal data layer is the monopoly.'])
slide(ff, 7, T, 'Market: wedge to monopoly',
      'A definite, large vision — TAM big enough for a power-law outcome.',
      ['Wedge: creator live commerce. [Insert live-commerce / creator-economy TAM with source.]',
       'Expansion: every video platform needs shoppable + ad attribution + accessibility — peekuu is the layer.',
       'End state: the commerce-and-attention graph for live video in the West.'])
slide(ff, 8, T, 'Business model — many layers, one viewer action',
      'Multiple revenue layers compounding on a single action.',
      ['Ads (CPM; $25 CPM floor in current model) + affiliate clicks on detected products.',
       'Gifting economy: coins -> creator diamonds -> cash-out (platform take rate); sponsored brand gifts.',
       'Brand-intelligence SaaS (the second act).'])
slide(ff, 9, T, 'Traction — built, fast',
      'Velocity + that the hard tech already works. Founders Fund funds builders.',
      ['Shipped in weeks: mobile (iOS/Android) + web studio + Firebase backend + Agora RTC.',
       'Live: multimodal detection, gifting + payouts, co-host, virtual-background rooms, post-stream editor, record-to-feed.',
       '[Insert any early usage / pilot creators / waitlist — or state pre-launch honestly.]'])
slide(ff, 10, T, 'Team',
      'Why YOU are the inevitable founder for this secret.',
      ['[Founder background: why you see this secret others do not.]',
       '[Unfair advantages: domain, distribution, technical depth.]'])
slide(ff, 11, T, 'The ask',
      'Specific raise + the singular milestone it buys.',
      ['[Raising $X to reach Y — e.g., live-commerce GMV / retained creators in N months.]',
       'Use of funds: [detection accuracy, creator acquisition, advertiser pilots].'])
ff.save('E:/claude/qaramia-v2/Peekuu-Pitch-FoundersFund.docx')

# ── BENCHMARK CUT — loop / retention / network first ──────────────────────────
bm = new_doc()
cover(bm, 'Benchmark', 'Engagement loop -> cohort retention -> network effects (Sarah Tavel frameworks). Metrics-led.')
T = 12
slide(bm, 1, T, 'The one-liner',
      'A consumer network with a clear, escalating engagement loop — not a feature.',
      ['peekuu is a live + short-video commerce network where watching turns into chatting, gifting, and buying.',
       'Every product shown or spoken becomes shoppable in real time — the loop monetizes itself.'], first=True)
slide(bm, 2, T, 'The engagement loop (hierarchy of engagement)',
      'Escalating engagement + accruing benefits + rising switching costs.',
      ['Watch -> danmaku comment -> send gift -> tip earned coins -> creator earns -> creator streams more -> richer feed.',
       'Accruing benefits: coin balance, follow graph, creator earnings + audience.',
       'Rising switching costs the more a user engages — on both sides.'])
slide(bm, 3, T, 'Retention is the thesis',
      'A cohort retention curve that FLATTENS (the smile). This is the slide Benchmark buys on.',
      ['[Insert D1 / D7 / D30 cohort retention curve — the flattening tail is the whole pitch.]',
       '[Viewer retention AND creator retention, shown separately.]',
       'If the curve is not flat yet, that is the #1 thing to prove before this meeting.'])
slide(bm, 4, T, 'The magic moment (live demo)',
      'Show the loop firing, and the commerce hook landing, in seconds.',
      ['Streamer says "I love this Stanley cup" -> shoppable card + matched ad appears for every viewer in ~2s.',
       'Viewer gifts -> creator earns -> shows up on the live leaderboard / gift goal.',
       'Demonstrates engagement and monetization in one motion.'])
slide(bm, 5, T, 'Why it is a network, not a feature',
      'Defensibility vs. a giant bolting on a clone — the core Benchmark objection.',
      ['Two-sided: creators bring audiences; advertisers bring demand; product catalog + advertiser graph compound.',
       'Multimodal detection improves with usage — a data network effect, not just UI.',
       'A clone cannot replicate creator earnings history, audience, or the detection corpus.'])
slide(bm, 6, T, 'The white space — 20x gap, no fused competitor',
      'The market is proven in Asia and structurally open in the West. (Figures fact-checked.)',
      ['Asia ~$370B live commerce vs. US ~$17B — proven demand, not a hypothesis.',
       'Feed-first beats store-first: Douyin took 47% of China live-commerce GMV vs. marketplace-native Taobao trailing.',
       'No Western app fuses feed + live + danmaku + auto-shopping + gifting — TikTok, Twitch, Whatnot, YouTube each own one fragment.',
       'TikTok pays creators as little as ~40% on gifts vs. Twitch/YouTube ~70% — peekuu\'s tiered split is a direct creator-acquisition wedge.'])
slide(bm, 7, T, 'Unit economics of one viewer action',
      'That engagement converts to durable, expanding margin per user.',
      ['Gifting: coins -> diamonds -> cash-out, with platform take rate [insert %].',
       'Ads ($25 CPM floor) + affiliate clicks on detected products; sponsored brand gifts.',
       '[Insert ARPDAU / contribution margin per active stream.]'])
slide(bm, 7, T, 'The detection moat (brief)',
      'The proprietary layer that makes the network compound — kept short here.',
      ['Barcode + Claude Vision + speech detection turns content into ad inventory automatically.',
       'Accuracy compounds with the catalog + transcript corpus -> a widening data moat.'])
slide(bm, 8, T, 'Market',
      'Large enough for a venture-scale network outcome.',
      ['Wedge: creator live commerce. [Insert TAM/SAM with source.]',
       'Adjacent: short-video commerce + creator monetization + brand intelligence.'])
slide(bm, 9, T, 'Traction & metrics',
      'Early signal of product-market fit — Benchmark wants proof, not promises.',
      ['[Core funnel: DAU/WAU, sessions/day, stream length, gift conversion, GMV.]',
       '[Creator metrics: active creators, repeat-stream rate, earnings.]',
       'Shipped end-to-end on mobile + web (live, gifting, payouts, detection, editor).'])
slide(bm, 10, T, 'Team',
      'Why this team out-executes on a consumer network.',
      ['[Founder background + why you understand this loop.]',
       '[Shipping velocity as evidence — full stack live in weeks.]'])
slide(bm, 11, T, 'The ask',
      'Specific raise tied to the retention/network milestone it funds.',
      ['[Raising $X to reach Y retained creators / cohort-retention target in N months.]',
       'Use of funds: [creator growth, retention loops, detection accuracy].'])
bm.save('E:/claude/qaramia-v2/Peekuu-Pitch-Benchmark.docx')

print('Wrote Peekuu-Pitch-FoundersFund.docx and Peekuu-Pitch-Benchmark.docx')
